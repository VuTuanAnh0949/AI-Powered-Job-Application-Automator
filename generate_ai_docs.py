#!/usr/bin/env python3
"""
AI Resume and Cover Letter Generator
Supports both Gemini and Groq APIs with interactive or command-line input
"""
import os
import sys
import json
import argparse
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt

# Add project to path
sys.path.insert(0, str(Path(__file__).parent))

from dotenv import load_dotenv
load_dotenv("job_application_automation/.env")


def detect_available_provider():
    """Detect which AI provider is available."""
    provider = os.getenv("LLAMA_API_PROVIDER", "").lower()
    gemini_key = os.getenv("GEMINI_API_KEY", "")
    groq_key = os.getenv("GROQ_API_KEY", "")
    
    # If provider specified and has key
    if provider == "gemini" and gemini_key and gemini_key != "your_gemini_api_key_here":
        return "gemini"
    elif provider == "groq" and groq_key and groq_key != "your_groq_api_key_here":
        return "groq"
    
    # Auto-detect based on available keys
    if groq_key and groq_key != "your_groq_api_key_here":
        return "groq"
    elif gemini_key and gemini_key != "your_gemini_api_key_here":
        return "gemini"
    
    return None


def load_candidate_profile():
    """Load candidate profile from JSON."""
    profile_path = Path("job_application_automation/data/candidate_profile.json")
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def generate_with_ai(prompt, provider="groq"):
    """Generate content using specified AI provider."""
    if provider == "groq":
        from groq import Groq
        api_key = os.getenv("GROQ_API_KEY")
        client = Groq(api_key=api_key)
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2048,
            temperature=0.7
        )
        return response.choices[0].message.content
        
    elif provider == "gemini":
        from google import genai
        api_key = os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key)
        
        # Try available models with new API (full paths required)
        models_to_try = [
            'models/gemini-2.5-flash',
            'models/gemini-2.5-pro',
            'models/gemini-2.0-flash',
            'models/gemini-flash-latest',
            'models/gemini-pro-latest'
        ]
        
        for model_name in models_to_try:
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                return response.text
            except Exception as e:
                print(f"Model {model_name} failed: {e}")
                continue
        raise Exception("No Gemini model available")
    
    raise Exception(f"Unknown provider: {provider}")


def generate_resume(candidate, job_description="Python Developer position", provider="groq"):
    """Generate resume using AI."""
    prompt = f"""
Create a professional resume based on this candidate information:

Name: {candidate['full_name']}
Email: {candidate['email']}
Phone: {candidate['phone']}
LinkedIn: {candidate.get('linkedin_url', '')}

Summary: {candidate['summary']}

Skills: {', '.join(candidate['skills'])}

Experience:
{chr(10).join([f"- {exp['title']} at {exp['company']} ({exp['dates']}): {exp['description']}" for exp in candidate.get('experience', [])])}

Target Position: {job_description}

Generate a well-formatted resume with these sections:
1. Professional Summary (concise, impactful)
2. Core Skills (organized by category)
3. Work Experience (highlight achievements with metrics)
4. Education (if mentioned)

Make it ATS-friendly, professional, and compelling. Focus on quantifiable achievements.
"""
    
    return generate_with_ai(prompt, provider)


def generate_cover_letter(candidate, company_name="the company", job_title="the position", 
                         job_description="", provider="groq"):
    """Generate cover letter using AI."""
    job_desc_text = f"\n\nJob Description:\n{job_description}" if job_description else ""
    
    prompt = f"""
Create a professional cover letter for:

Candidate: {candidate['full_name']}
Applying for: {job_title}
Company: {company_name}

Candidate Background: {candidate['summary']}

Key Skills: {', '.join(candidate['skills'][:10])}

Recent Experience:
{chr(10).join([f"- {exp['title']} at {exp['company']}: {exp['description'][:150]}..." for exp in candidate.get('experience', [])[:2]])}
{job_desc_text}

Write a compelling cover letter that:
1. Opens with a strong hook mentioning the specific position
2. Highlights 2-3 most relevant experiences/achievements with specific metrics
3. Shows genuine enthusiasm for the company and role
4. Demonstrates cultural fit and understanding of the company
5. Ends with a strong call to action

Keep it concise (350-450 words). Professional yet personable tone. Use specific examples.
"""
    
    return generate_with_ai(prompt, provider)


def save_to_docx(content, filename, title="Document"):
    """Save content to a Word document."""
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Add content - split by lines
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            if para.startswith('##'):
                # Heading level 2
                doc.add_heading(para.strip('# '), level=2)
            elif para.startswith('#'):
                # Heading level 1
                doc.add_heading(para.strip('# '), level=1)
            else:
                # Normal paragraph
                p = doc.add_paragraph(para)
                p.style.font.size = Pt(11)
    
    doc.save(filename)
    print(f"   ✅ {filename.name}")


def get_user_input():
    """Get job details from user interactively."""
    print("\n" + "="*60)
    print("📝 Enter Job Details (press Enter to use defaults)")
    print("="*60)
    
    company = input("\n🏢 Company Name [Your Target Company]: ").strip()
    if not company:
        company = "Your Target Company"
    
    job_title = input("💼 Job Title [AI Engineer / Data Scientist]: ").strip()
    if not job_title:
        job_title = "AI Engineer / Data Scientist"
    
    print("\n📄 Job Description (optional - paste and press Enter twice):")
    print("   [Leave empty for general application]")
    lines = []
    empty_count = 0
    while empty_count < 2:
        line = input()
        if line:
            lines.append(line)
            empty_count = 0
        else:
            empty_count += 1
            if empty_count < 2:
                lines.append('')
    
    job_description = '\n'.join(lines).strip()
    
    return company, job_title, job_description


def main():
    parser = argparse.ArgumentParser(
        description='Generate AI-powered Resume and Cover Letter',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Interactive mode (will prompt for details)
  python generate_ai_docs.py
  
  # Quick generation with command-line args
  python generate_ai_docs.py --company "Google" --job "Software Engineer"
  
  # With full job description
  python generate_ai_docs.py --company "Microsoft" --job "Data Scientist" --description "Looking for..."
        """
    )
    parser.add_argument('--company', type=str, help='Company name')
    parser.add_argument('--job', type=str, help='Job title')
    parser.add_argument('--description', type=str, help='Job description')
    parser.add_argument('--interactive', '-i', action='store_true', 
                       help='Force interactive mode')
    
    args = parser.parse_args()
    
    print("\n🚀 AI Resume & Cover Letter Generator")
    print("="*60)
    
    try:
        # Detect provider
        provider = detect_available_provider()
        if not provider:
            print("\n❌ Error: No AI provider available!")
            print("\n⚠️  Please set up one of these in .env file:")
            print("   1. GROQ_API_KEY (recommended - get at: https://console.groq.com/keys)")
            print("   2. GEMINI_API_KEY (get at: https://aistudio.google.com/app/apikey)")
            return
        
        print(f"🤖 Using AI Provider: {provider.upper()}")
        
        # Load candidate info
        print("\n📋 Loading candidate profile...")
        candidate = load_candidate_profile()
        print(f"   ✅ Loaded: {candidate['full_name']}")
        
        # Get job details
        if args.interactive or (not args.company and not args.job):
            company, job_title, job_description = get_user_input()
        else:
            company = args.company or "Your Target Company"
            job_title = args.job or "AI Engineer / Data Scientist"
            job_description = args.description or ""
        
        print(f"\n🎯 Target: {job_title} at {company}")
        
        # Generate documents
        print("\n📝 Generating documents with AI...")
        print("   ⏳ Creating Resume...")
        resume_text = generate_resume(
            candidate, 
            job_description=job_title if not job_description else f"{job_title}\n\n{job_description}",
            provider=provider
        )
        
        print("   ⏳ Creating Cover Letter...")
        cover_letter_text = generate_cover_letter(
            candidate,
            company_name=company,
            job_title=job_title,
            job_description=job_description,
            provider=provider
        )
        
        # Save files
        output_dir = Path("job_application_automation/data/generated_cover_letters")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company = "".join(c for c in company if c.isalnum() or c in (' ', '-', '_'))[:30]
        
        resume_file = output_dir / f"{timestamp}_{safe_company}_resume.docx"
        cover_letter_file = output_dir / f"{timestamp}_{safe_company}_cover_letter.docx"
        
        print("\n💾 Saving documents...")
        save_to_docx(resume_text, resume_file, f"Resume - {candidate['full_name']}")
        save_to_docx(cover_letter_text, cover_letter_file, 
                    f"Cover Letter - {job_title} at {company}")
        
        print("\n" + "="*60)
        print("✅ AI Generation Complete!")
        print("="*60)
        print(f"\n📂 Location: {output_dir}")
        print(f"\n📄 Files created:")
        print(f"   • {resume_file.name}")
        print(f"   • {cover_letter_file.name}")
        print("\n💡 Tip: Review and customize documents before sending!")
        
    except KeyboardInterrupt:
        print("\n\n❌ Cancelled by user")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\n💡 Make sure:")
        print("   1. API key is set in .env file")
        print("   2. candidate_profile.json exists and is valid")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
