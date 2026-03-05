#!/usr/bin/env python3
"""
Simple Resume and Cover Letter Generator using Groq API
"""
import os
import sys
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt

# Add project to path
sys.path.insert(0, str(Path(__file__).parent))

from dotenv import load_dotenv
load_dotenv("job_application_automation/.env")

from groq import Groq

def load_candidate_profile():
    """Load candidate profile from JSON."""
    profile_path = Path("job_application_automation/data/candidate_profile.json")
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def generate_resume(candidate, job_description="Python Developer position"):
    """Generate resume using Groq."""
    api_key = os.getenv("GROQ_API_KEY")
    client = Groq(api_key=api_key)
    
    prompt = f"""
    Create a professional resume based on this candidate information:
    
    Name: {candidate['full_name']}
    Email: {candidate['email']}
    Phone: {candidate['phone']}
    LinkedIn: {candidate.get('linkedin_url', '')}
    
    Summary: {candidate['summary']}
    
    Skills: {', '.join(candidate['skills'])}
    
    Experience:
    {chr(10).join([f"- {exp['title']} at {exp['company']} ({exp['dates']}): {exp['description']}" for exp in candidate.get('experience', [])])}
    
    Target Position: {job_description}
    
    Generate a well-formatted resume text with sections: Professional Summary, Core Skills, Work Experience, Education.
    Make it ATS-friendly and professional. Maximum 1 page worth of content.
    """
    
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
        temperature=0.7
    )
    
    return response.choices[0].message.content

def generate_cover_letter(candidate, company_name="the company", job_title="the position"):
    """Generate cover letter using Groq."""
    api_key = os.getenv("GROQ_API_KEY")
    client = Groq(api_key=api_key)
    
    prompt = f"""
    Create a professional cover letter for:
    
    Candidate: {candidate['full_name']}
    Applying for: {job_title}
    Company: {company_name}
    
    Background: {candidate['summary']}
    
    Key Skills: {', '.join(candidate['skills'][:10])}
    
    Write a compelling cover letter that:
    1. Opens professionally
    2. Highlights relevant experience and skills
    3. Shows enthusiasm for the role
    4. Ends with a strong call to action
    
    Keep it concise (300-400 words). Professional but personable tone.
    """
    
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
        temperature=0.7
    )
    
    return response.choices[0].message.content

def save_to_docx(content, filename, title="Document"):
    """Save content to a Word document."""
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Add content - split by lines
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            if para.startswith('#'):
                # Heading
                doc.add_heading(para.strip('# '), level=1)
            elif para.startswith('##'):
                doc.add_heading(para.strip('## '), level=2)
            else:
                # Normal paragraph
                p = doc.add_paragraph(para)
                p.style.font.size = Pt(11)
    
    doc.save(filename)
    print(f"✅ Saved: {filename}")

def main():
    print("🚀 AI Resume & Cover Letter Generator (Groq)\n")
    print("="*60)
    
    try:
        # Load candidate info
        print("\n📋 Loading candidate profile...")
        candidate = load_candidate_profile()
        print(f"   Loaded profile for: {candidate['full_name']}")
        
        # Generate resume
        print("\n📝 Generating Resume with AI...")
        resume_text = generate_resume(candidate)
        
        # Save resume
        output_dir = Path("job_application_automation/data/generated_cover_letters")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        resume_file = output_dir / f"{timestamp}_ai_resume.docx"
        save_to_docx(resume_text, resume_file, f"Resume - {candidate['full_name']}")
        
        # Generate cover letter
        print("\n💌 Generating Cover Letter with AI...")
        cover_letter_text = generate_cover_letter(
            candidate,
            company_name="Your Target Company",
            job_title="AI Engineer / Data Scientist"
        )
        
        # Save cover letter
        cover_letter_file = output_dir / f"{timestamp}_ai_cover_letter.docx"
        save_to_docx(cover_letter_text, cover_letter_file, "Cover Letter")
        
        print("\n" + "="*60)
        print("✅ AI Generation Complete!")
        print("="*60)
        print(f"\n📁 Files saved to:")
        print(f"   {resume_file}")
        print(f"   {cover_letter_file}")
        
        print("\n💡 Documents generated with AI - review and customize before sending!")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\n⚠️  Make sure you have:")
        print("   1. GROQ_API_KEY set in .env file")
        print("   2. Get free key at: https://console.groq.com/keys")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
