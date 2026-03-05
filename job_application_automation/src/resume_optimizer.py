"""
ATS (Applicant Tracking System) scoring and resume optimization module.
This module provides functionality to score and optimize resumes against job descriptions.
"""

import os
import re
import json
import logging
import numpy as np
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Set, Union
from sentence_transformers import SentenceTransformer
import docx
from docx import Document
import faiss
from textblob import TextBlob
import spacy
import pandas as pd
import requests

# Import project modules
from job_application_automation.src.resume_cover_letter_generator import ResumeGenerator
from job_application_automation.config.llama_config import LlamaConfig
from job_application_automation.src.utils.path_utils import get_project_root, get_data_path

# Set up logging
_data_dir = get_data_path()
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(str(_data_dir / "resume_optimizer.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Try to load spaCy model for NLP tasks
try:
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    logger.warning("spaCy model not available. Some NLP features will be limited.")
    SPACY_AVAILABLE = False

# Define paths
DATA_DIR = _data_dir
ATS_INDEX_PATH = DATA_DIR / "ats_index.idx"
ATS_META_PATH = DATA_DIR / "ats_meta.json"

# ATS score threshold for "good" resume
ATS_THRESHOLD = 0.75


class ATSScorer:
    """
    Class for scoring resumes against job descriptions using ATS-like algorithms.
    """

    def __init__(self):
        """Initialize the ATSScorer with embedding model."""
        os.makedirs(DATA_DIR, exist_ok=True)
        
        # Initialize embedding model with proper fallback mechanisms
        self.embedding_model = None
        self.embedding_dim = 384  # Default dimension for all-MiniLM-L6-v2
        
        try:
            # First try to load the model
            logger.info("Loading sentence transformer model...")
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
            logger.info(f"Successfully loaded sentence transformer model with dimension {self.embedding_dim}")
        except (OSError, ConnectionError, requests.exceptions.ConnectionError) as e:
            logger.warning(f"Failed to load sentence transformer model: {e}. Using fallback embedding method.")
            
            # Try to load a local model if available
            try:
                local_model_path = os.path.join(os.path.expanduser('~'), '.cache', 'sentence_transformers')
                if os.path.exists(local_model_path):
                    logger.info(f"Attempting to load local model from {local_model_path}")
                    self.embedding_model = SentenceTransformer(local_model_path)
                    self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
                    logger.info("Successfully loaded local sentence transformer model")
                else:
                    logger.warning("No local model found. Using fallback text comparison.")
            except Exception as local_error:
                logger.error(f"Failed to load local model: {local_error}")
        
        # Set flag for embedding availability
        if self.embedding_model is None:
            logger.error("Failed to load embedding model - semantic features disabled")
            self.embedding_disabled = True
        else:
            self.embedding_disabled = False
                
        # Initialize FAISS index for resume-job matching
        self._setup_vector_index()
        
        # Dictionary to store keyword importance by industry/role
        self.industry_keywords = self._load_industry_keywords()
        
    def _setup_vector_index(self) -> None:
        """Set up FAISS vector index for resume embeddings."""
        try:
            if ATS_INDEX_PATH.exists() and ATS_META_PATH.exists():
                # Load existing index
                self.index = faiss.read_index(str(ATS_INDEX_PATH))
                with open(ATS_META_PATH, 'r') as f:
                    self.meta = json.load(f)
                # Determine next id
                self.next_id = max((item['id'] for item in self.meta), default=0) + 1
            else:
                # Create new index
                base_index = faiss.IndexFlatL2(self.embedding_dim)
                self.index = faiss.IndexIDMap(base_index)
                self.meta = []
                self.next_id = 0
                
            logger.info("FAISS vector index initialized for ATS scoring")
        except Exception as e:
            logger.error(f"Error setting up FAISS index: {e}")
            # Fallback to empty index
            base_index = faiss.IndexFlatL2(self.embedding_dim)
            self.index = faiss.IndexIDMap(base_index)
            self.meta = []
            self.next_id = 0
    
    def _load_industry_keywords(self) -> Dict[str, Dict[str, float]]:
        """Load industry-specific keywords with weights."""
        # Default weights for common roles
        industry_keywords = {
            "software_development": {
                "python": 0.8, "java": 0.7, "javascript": 0.7, "react": 0.6,
                "nodejs": 0.6, "aws": 0.7, "docker": 0.6, "kubernetes": 0.6,
                "microservices": 0.6, "ci/cd": 0.6, "git": 0.5, "agile": 0.5,
                "testing": 0.5, "rest api": 0.6, "sql": 0.6
            },
            "data_science": {
                "python": 0.9, "r": 0.7, "sql": 0.7, "machine learning": 0.9,
                "deep learning": 0.8, "tensorflow": 0.7, "pytorch": 0.7,
                "pandas": 0.7, "numpy": 0.7, "scikit-learn": 0.7,
                "data visualization": 0.6, "statistics": 0.7, "nlp": 0.6,
                "computer vision": 0.6, "big data": 0.6, "spark": 0.6
            },
            "product_management": {
                "product development": 0.8, "agile": 0.7, "scrum": 0.6,
                "roadmap": 0.7, "user research": 0.7, "product strategy": 0.8,
                "stakeholder management": 0.7, "kpi": 0.7, "analytics": 0.6,
                "a/b testing": 0.6, "product lifecycle": 0.7
            },
            "marketing": {
                "digital marketing": 0.8, "content marketing": 0.7, "seo": 0.7,
                "sem": 0.6, "social media": 0.7, "marketing strategy": 0.8,
                "google analytics": 0.7, "conversion rate": 0.6, "crm": 0.6,
                "email marketing": 0.6, "campaign management": 0.7
            }
        }
        
        # Try to load custom keywords if available
        keywords_path = DATA_DIR / "industry_keywords.json"
        if keywords_path.exists():
            try:
                with open(keywords_path, 'r') as f:
                    custom_keywords = json.load(f)
                # Merge with defaults
                for industry, keywords in custom_keywords.items():
                    if industry in industry_keywords:
                        industry_keywords[industry].update(keywords)
                    else:
                        industry_keywords[industry] = keywords
                logger.info("Loaded custom industry keywords")
            except Exception as e:
                logger.warning(f"Error loading industry keywords: {e}")
        
        return industry_keywords
    
    def save_index(self) -> None:
        """Save the FAISS index to disk."""
        try:
            faiss.write_index(self.index, str(ATS_INDEX_PATH))
            with open(ATS_META_PATH, 'w') as f:
                json.dump(self.meta, f, indent=2)
            logger.info("Saved ATS index to disk")
        except Exception as e:
            logger.error(f"Error saving ATS index: {e}")
    
    def parse_resume(self, resume_path: str) -> Dict[str, Any]:
        """
        Parse resume from file.
        
        Args:
            resume_path: Path to the resume file (.docx, .pdf, .txt, or .md)
            
        Returns:
            Dictionary with parsed resume sections
        """
        resume_text = self._read_document(resume_path)
        if not resume_text:
            logger.error(f"Could not read resume from {resume_path}")
            return {}
        
        # Parse resume into sections
        sections = self._parse_resume_sections(resume_text, os.path.splitext(resume_path)[1])
        
        # Extract additional structured data
        sections["keywords"] = self._extract_keywords(resume_text)
        sections["file_path"] = resume_path
        
        logger.info(f"Successfully parsed resume from {resume_path}")
        return sections
    
    def _read_document(self, file_path: str) -> str:
        """Read document content from file."""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.docx':
                doc = Document(file_path)
                return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ' '.join(page.extract_text() for page in reader.pages)
                        return text
                except ImportError:
                    logger.warning("PyPDF2 not installed. Cannot read PDF files.")
                    return ""
            elif ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:  # Default to text file
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return ""
    
    def _parse_resume_sections(self, text: str, file_ext: str) -> Dict[str, Any]:
        """Parse resume text into structured sections."""
        sections = {}
        
        # Basic sections to look for
        section_patterns = {
            "contact": [r"contact\s+information", r"email|phone|address", r"^.{0,50}@.+\..{2,}"],
            "summary": [r"summary|profile|objective", r"professional\s+summary", r"career\s+objective"],
            "experience": [r"experience|employment|work\s+history", r"professional\s+experience"],
            "education": [r"education|academic|degree", r"university|college|school"],
            "skills": [r"skills|expertise|technologies", r"technical\s+skills", r"competencies"],
            "projects": [r"projects|portfolio", r"key\s+projects", r"project\s+experience"],
            "certifications": [r"certifications|certificates", r"professional\s+development"],
        }
        
        if file_ext.lower() == '.md':
            # For markdown files, look for headings
            sections = self._parse_markdown_resume(text)
        else:
            # For other formats, use regex pattern matching
            current_section = "header"
            lines = text.split('\n')
            sections[current_section] = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if line starts a new section
                new_section_detected = False
                for section_name, patterns in section_patterns.items():
                    for pattern in patterns:
                        if re.search(pattern, line, re.IGNORECASE) and len(line) < 100:
                            current_section = section_name
                            sections[current_section] = [line]
                            new_section_detected = True
                            break
                    if new_section_detected:
                        break
                
                if not new_section_detected:
                    if current_section in sections:
                        sections[current_section].append(line)
                    else:
                        sections[current_section] = [line]
            
            # Join section lines
            for section, lines in sections.items():
                sections[section] = "\n".join(lines)
        
        # Extract contact information
        if "contact" in sections or "header" in sections:
            contact_text = sections.get("contact", sections.get("header", ""))
            sections["email"] = self._extract_email(contact_text)
            sections["phone"] = self._extract_phone(contact_text)
        
        # Extract name
        sections["name"] = self._extract_name(sections.get("header", ""))
        
        return sections
    
    def _parse_markdown_resume(self, text: str) -> Dict[str, Any]:
        """Parse a markdown resume into sections."""
        sections = {}
        current_section = "header"
        section_content = []
        
        # Split by lines and process
        lines = text.split('\n')
        for line in lines:
            # Check if line is a markdown heading
            if line.strip().startswith('# '):
                # Save previous section
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                
                # Start new section
                heading = line.strip().replace('# ', '').lower()
                # Clean up section name
                section_name = re.sub(r'[^a-z0-9_\s]', '', heading).strip().replace(' ', '_')
                if not section_name:
                    section_name = "section_" + str(len(sections))
                current_section = section_name
                section_content = []
            else:
                section_content.append(line)
        
        # Save last section
        if section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections
    
    def _extract_email(self, text: str) -> str:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group(0) if match else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number from text."""
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        match = re.search(phone_pattern, text)
        return match.group(0) if match else ""
    
    def _extract_name(self, text: str) -> str:
        """Extract name from header section."""
        lines = text.strip().split('\n')
        if lines:
            # Usually the first non-empty line is the name
            for line in lines:
                if line.strip() and len(line.strip()) < 50:
                    # Remove common markdown formatting
                    name = re.sub(r'[#*_]', '', line.strip())
                    return name
        return ""
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract keywords from text."""
        keywords = []
        
        if SPACY_AVAILABLE:
            doc = nlp(text)
            # Extract nouns and proper nouns
            keywords = [token.lemma_ for token in doc if token.pos_ in ('NOUN', 'PROPN')]
            
            # Extract noun phrases
            noun_phrases = [chunk.text for chunk in doc.noun_chunks]
            keywords.extend(noun_phrases)
        else:
            # Fallback to simple text processing
            blob = TextBlob(text)
            keywords = [word for word, tag in blob.tags if tag.startswith('NN')]
            
            # Simple noun phrase extraction
            words = text.lower().split()
            for i in range(len(words) - 1):
                if len(words[i]) > 3 and len(words[i+1]) > 3:
                    keywords.append(f"{words[i]} {words[i+1]}")
        
        # Clean up and deduplicate
        cleaned_keywords = []
        for kw in keywords:
            # Clean up and normalize
            cleaned = re.sub(r'[^\w\s]', '', kw).strip().lower()
            if cleaned and len(cleaned) > 2 and not cleaned.isdigit():
                cleaned_keywords.append(cleaned)
        
        return list(set(cleaned_keywords))
    
    def score_resume(self, resume: Dict[str, Any], job_description: str) -> Dict[str, Any]:
        """
        Score a parsed resume against a job description.
        
        Args:
            resume: Parsed resume dictionary
            job_description: Job description text
            
        Returns:
            Dictionary with ATS score and details
        """
        # Extract job requirements
        job_keywords = self._extract_job_requirements(job_description)
        job_role_type = self._detect_job_role_type(job_description)
        
        # Calculate different scoring components
        keyword_score = self._calculate_keyword_match(resume, job_keywords, job_role_type)
        format_score = self._check_resume_format(resume)
        semantic_score = self._calculate_semantic_similarity(resume, job_description)
        
        # Calculate final weighted score
        final_score = (
            keyword_score * 0.5 +   # Keywords are most important
            semantic_score * 0.3 +  # Semantic match is important
            format_score * 0.2      # Format matters less
        )
        
        # Get detailed analysis
        keyword_analysis = self._analyze_keyword_matches(resume, job_keywords, job_role_type)
        missing_keywords = self._identify_missing_keywords(resume, job_keywords, job_role_type)
        improvement_suggestions = self._generate_improvement_suggestions(resume, job_description, missing_keywords)
        
        # Create score report
        score_details = {
            "overall_score": round(final_score * 100, 1),  # Convert to percentage
            "keyword_score": round(keyword_score * 100, 1),
            "semantic_score": round(semantic_score * 100, 1),
            "format_score": round(format_score * 100, 1),
            "matched_keywords": keyword_analysis["matches"],
            "missing_important_keywords": missing_keywords,
            "improvement_suggestions": improvement_suggestions,
            "job_role_detected": job_role_type
        }
        
        # Add resume to index for future reference
        self._add_resume_to_index(resume, job_description, final_score)
        
        logger.info(f"Resume scored: {score_details['overall_score']}% overall match")
        return score_details
    
    def _extract_job_requirements(self, job_description: str) -> List[Dict[str, Any]]:
        """Extract key requirements and keywords from job description."""
        keywords = []
        
        # First try spaCy if available
        if SPACY_AVAILABLE:
            doc = nlp(job_description)
            
            # Extract skill-related noun phrases
            for chunk in doc.noun_chunks:
                if len(chunk.text) > 3:
                    keywords.append({
                        "keyword": chunk.text.lower(),
                        "weight": 0.5  # Default weight
                    })
            
            # Extract technology terms with higher weight
            tech_patterns = [r"Python", r"Java\b", r"JavaScript", r"React", r"AWS", r"SQL", 
                            r"Excel", r"Machine Learning", r"AI", r"R\b", r"C\+\+",
                            r"Docker", r"Kubernetes", r"Linux", r"Git"]
            
            for pattern in tech_patterns:
                matches = re.findall(pattern, job_description, re.IGNORECASE)
                for match in matches:
                    keywords.append({
                        "keyword": match.lower(),
                        "weight": 0.8  # Higher weight for tech terms
                    })
        
        # Look for "required" and "preferred" sections
        required_section = re.search(r"required(?:\s+qualifications|\s+skills)?:(.*?)(?:preferred|$)", 
                                    job_description.lower(), re.DOTALL)
        if required_section:
            # Extract bullet points or lines
            for line in required_section.group(1).split('\n'):
                line = line.strip()
                if line and (line.startswith('-') or line.startswith('•')):
                    keywords.append({
                        "keyword": line.strip('- •').strip().lower(),
                        "weight": 0.9  # Higher weight for explicitly required skills
                    })
        
        # Look for common job requirement patterns
        requirement_patterns = [
            r"(\d+)(?:\+)?\s+years\s+(?:of\s+)?experience\s+(?:in|with)\s+([^,.]+)",
            r"proficient\s+(?:in|with)\s+([^,.]+)",
            r"knowledge\s+of\s+([^,.]+)",
            r"familiarity\s+with\s+([^,.]+)",
            r"background\s+in\s+([^,.]+)",
            r"degree\s+in\s+([^,.]+)"
        ]
        
        for pattern in requirement_patterns:
            matches = re.finditer(pattern, job_description, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) > 1:
                    # This is a pattern with years of experience
                    years = match.group(1)
                    skill = match.group(2).lower()
                    keywords.append({
                        "keyword": skill,
                        "weight": 0.9,
                        "years_required": int(years)
                    })
                else:
                    keywords.append({
                        "keyword": match.group(1).lower(),
                        "weight": 0.7
                    })
        
        # Deduplicate keywords based on similarity
        unique_keywords = []
        seen_keywords = set()
        
        for item in keywords:
            keyword = item["keyword"]
            if keyword not in seen_keywords:
                seen_keywords.add(keyword)
                unique_keywords.append(item)
        
        return unique_keywords
    
    def _detect_job_role_type(self, job_description: str) -> str:
        """Detect job role category from job description."""
        job_description_lower = job_description.lower()
        
        # Define role patterns
        role_patterns = {
            "software_development": [
                "software engineer", "developer", "programmer", "coder", "full stack",
                "frontend", "backend", "web developer", "mobile developer"
            ],
            "data_science": [
                "data scientist", "machine learning", "AI engineer", "data analyst",
                "deep learning", "NLP", "computer vision", "statistical analysis"
            ],
            "product_management": [
                "product manager", "product owner", "product development",
                "roadmap", "user stories", "prioritization"
            ],
            "marketing": [
                "marketing", "social media", "content writer", "seo", "ppc",
                "growth hacker", "brand manager"
            ],
            "design": [
                "designer", "ui/ux", "graphic designer", "ui designer", "ux researcher"
            ],
            "sales": [
                "sales", "account executive", "business development", "sales representative"
            ],
            "finance": [
                "finance", "accountant", "financial analyst", "controller", "cfo"
            ],
            "hr": [
                "human resources", "recruiter", "talent acquisition", "hr manager"
            ],
            "operations": [
                "operations", "project manager", "program manager", "scrum master"
            ]
        }
        
        # Score each role type
        role_scores = {}
        
        for role, patterns in role_patterns.items():
            score = sum(1 for pattern in patterns if pattern in job_description_lower)
            role_scores[role] = score
        
        # Get the role with highest score
        if role_scores:
            best_role = max(role_scores.items(), key=lambda x: x[1])
            if best_role[1] > 0:
                return best_role[0]
        
        # Default to generic if no clear match
        return "generic"
    
    def _calculate_keyword_match(self, resume: Dict[str, Any], 
                               job_keywords: List[Dict[str, Any]],
                               job_role_type: str) -> float:
        """Calculate keyword match score."""
        if not job_keywords:
            return 0.5  # Default score if no keywords extracted
        
        # Get resume text and keywords
        resume_text = self._get_full_resume_text(resume)
        resume_text_lower = resume_text.lower()
        
        # Track matches and their weights
        total_weight = 0.0
        matched_weight = 0.0
        
        # Check for presence of each job keyword in resume
        for keyword_item in job_keywords:
            keyword = keyword_item["keyword"].lower()
            weight = keyword_item["weight"]
            total_weight += weight
            
            # Check for exact or partial matches
            if keyword in resume_text_lower:
                matched_weight += weight
            elif len(keyword.split()) > 1:
                # For multi-word keywords, check if most words appear
                words = keyword.split()
                matched_words = sum(1 for word in words if word in resume_text_lower)
                if matched_words / len(words) >= 0.5:
                    matched_weight += weight * 0.7  # Partial credit
            
        # Check for role-specific keywords if available
        if job_role_type in self.industry_keywords:
            role_keywords = self.industry_keywords[job_role_type]
            role_total_weight = 0.0
            role_matched_weight = 0.0
            
            for keyword, weight in role_keywords.items():
                role_total_weight += weight
                if keyword.lower() in resume_text_lower:
                    role_matched_weight += weight
            
            # Only include if we have role keywords
            if role_total_weight > 0:
                # Add role-specific score (with lower weight)
                total_weight += role_total_weight * 0.5
                matched_weight += role_matched_weight * 0.5
        
        # Calculate final score
        if total_weight > 0:
            return matched_weight / total_weight
        else:
            return 0.5  # Default score
    
    def _get_full_resume_text(self, resume: Dict[str, Any]) -> str:
        """Get full text from resume sections."""
        text_parts = []
        
        for key, value in resume.items():
            if isinstance(value, str) and key not in ["file_path", "name", "email", "phone"]:
                text_parts.append(value)
        
        return " ".join(text_parts)
    
    def _check_resume_format(self, resume: Dict[str, Any]) -> float:
        """Check resume format quality."""
        score = 0.5  # Start with neutral score
        
        # Check presence of key sections
        important_sections = ["summary", "experience", "education", "skills"]
        section_score = sum(1 for section in important_sections if any(section in key.lower() for key in resume))
        section_score = min(section_score / len(important_sections), 1.0)
        
        # Check contact information
        has_email = bool(resume.get("email", ""))
        has_phone = bool(resume.get("phone", ""))
        has_name = bool(resume.get("name", ""))
        contact_score = (has_email + has_phone + has_name) / 3
        
        # Combine scores with weights
        score = section_score * 0.7 + contact_score * 0.3
        
        return score
    
    def _calculate_semantic_similarity(self, resume: Dict[str, Any], job_description: str) -> float:
        """Calculate semantic similarity between resume and job description."""
        resume_text = self._get_full_resume_text(resume)
        
        # Check if embedding model is available
        if self.embedding_model is None:
            # Fallback to simple word overlap calculation if embedding model not available
            resume_words = set(resume_text.lower().split())
            job_words = set(job_description.lower().split())
            
            if not resume_words or not job_words:
                return 0.5  # Default similarity
                
            # Calculate Jaccard similarity (intersection over union)
            intersection = len(resume_words.intersection(job_words))
            union = len(resume_words.union(job_words))
            
            if union == 0:
                return 0.5
            return float(intersection / union)
        
        # Generate embeddings with model if available
        try:
            if self.embedding_disabled or self.embedding_model is None:
                logger.warning("Embedding model not available, using keyword-only scoring")
                # Fallback to keyword-based scoring
                resume_keywords = set(resume_text.lower().split())
                job_keywords = set(job_description.lower().split())
                intersection = resume_keywords.intersection(job_keywords)
                union = resume_keywords.union(job_keywords)
                if not union:
                    return 0.5
                return float(len(intersection) / len(union))
            
            resume_embedding = self.embedding_model.encode(resume_text)
            job_embedding = self.embedding_model.encode(job_description)
            
            # Calculate cosine similarity
            resume_embedding = resume_embedding / np.linalg.norm(resume_embedding)
            job_embedding = job_embedding / np.linalg.norm(job_embedding)
            similarity = np.dot(resume_embedding, job_embedding)
            
            return float(similarity)
        except Exception as e:
            logger.error(f"Error calculating semantic similarity: {e}")
            return 0.5  # Default similarity
    
    def _analyze_keyword_matches(self, resume: Dict[str, Any], 
                              job_keywords: List[Dict[str, Any]], 
                              job_role_type: str) -> Dict[str, Any]:
        """Analyze which keywords from job description match in resume."""
        resume_text = self._get_full_resume_text(resume).lower()
        
        # Track matches
        matches = []
        misses = []
        
        # Check job specific keywords
        for keyword_item in job_keywords:
            keyword = keyword_item["keyword"].lower()
            if keyword in resume_text:
                matches.append({
                    "keyword": keyword,
                    "weight": keyword_item["weight"]
                })
            else:
                # Check for partial matches in multi-word keywords
                if len(keyword.split()) > 1:
                    words = keyword.split()
                    matched_words = sum(1 for word in words if word in resume_text)
                    if matched_words / len(words) >= 0.5:
                        matches.append({
                            "keyword": keyword,
                            "weight": keyword_item["weight"] * 0.7,
                            "partial": True
                        })
                    else:
                        misses.append({
                            "keyword": keyword,
                            "weight": keyword_item["weight"]
                        })
                else:
                    misses.append({
                        "keyword": keyword,
                        "weight": keyword_item["weight"]
                    })
        
        # Check role-specific keywords
        if job_role_type in self.industry_keywords:
            role_keywords = self.industry_keywords[job_role_type]
            for keyword, weight in role_keywords.items():
                if keyword.lower() in resume_text:
                    matches.append({
                        "keyword": keyword,
                        "weight": weight * 0.5,
                        "industry": True
                    })
                else:
                    misses.append({
                        "keyword": keyword,
                        "weight": weight * 0.5,
                        "industry": True
                    })
        
        # Sort by weight
        matches.sort(key=lambda x: x["weight"], reverse=True)
        misses.sort(key=lambda x: x["weight"], reverse=True)
        
        return {
            "matches": matches,
            "misses": misses
        }
    
    def _identify_missing_keywords(self, resume: Dict[str, Any], 
                                job_keywords: List[Dict[str, Any]],
                                job_role_type: str) -> List[Dict[str, Any]]:
        """Identify important keywords missing from resume."""
        analysis = self._analyze_keyword_matches(resume, job_keywords, job_role_type)
        
        # Filter for important missing keywords (high weight)
        important_missing = []
        for item in analysis["misses"]:
            if item["weight"] >= 0.7:  # Only include high-importance keywords
                important_missing.append(item)
            elif item.get("industry") and item["weight"] >= 0.4:
                # Include industry-specific keywords of medium importance
                important_missing.append(item)
        
        # Limit to most important
        return important_missing[:10]
    
    def _generate_improvement_suggestions(self, resume: Dict[str, Any],
                                       job_description: str,
                                       missing_keywords: List[Dict[str, Any]]) -> List[str]:
        """Generate suggestions for resume improvement."""
        suggestions = []
        
        # Suggest adding missing important keywords
        if missing_keywords:
            keywords_to_add = [item["keyword"] for item in missing_keywords[:5]]
            if keywords_to_add:
                suggestions.append(f"Add these missing keywords: {', '.join(keywords_to_add)}")
        
        # Check for formatting issues
        important_sections = ["summary", "experience", "education", "skills"]
        missing_sections = [section for section in important_sections 
                          if not any(section in key.lower() for key in resume)]
        
        if missing_sections:
            suggestions.append(f"Add these missing sections: {', '.join(missing_sections)}")
        
        # Check for contact information
        if not resume.get("email"):
            suggestions.append("Add your email address")
        if not resume.get("phone"):
            suggestions.append("Add your phone number")
            
        # Check experience descriptions
        if "experience" in resume:
            exp_text = resume["experience"].lower()
            if not re.search(r"achiev|result|impact|improve|increase|decrease|develop|create|launch|manage", exp_text):
                suggestions.append("Add more achievements and results in your experience section")
        
        return suggestions
    
    def _add_resume_to_index(self, resume: Dict[str, Any], job_description: str, score: float) -> None:
        """Add resume to vector index for future reference."""
        try:
            resume_text = self._get_full_resume_text(resume)
            
            # Create combined text for embedding
            combined_text = f"{resume_text}\n{job_description}"
            
            # Check if embedding model is available
            if self.embedding_disabled or self.embedding_model is None:
                logger.warning("Embedding model not available, skipping index addition")
                return
            
            # Generate embedding
            embedding = self.embedding_model.encode(combined_text)
            vec = embedding.reshape(1, -1).astype('float32')
            
            # Add to index
            self.index.add_with_ids(vec, np.array([self.next_id], dtype='int64'))
            
            # Store metadata
            meta_item = {
                "id": self.next_id,
                "resume_path": resume.get("file_path", ""),
                "score": score,
                "timestamp": datetime.now().isoformat()
            }
            self.meta.append(meta_item)
            self.next_id += 1
            
            # Save periodically
            if len(self.meta) % 10 == 0:
                self.save_index()
                
        except Exception as e:
            logger.error(f"Error adding resume to index: {e}")
    
    def find_similar_resumes(self, job_description: str, k: int = 5) -> List[Dict[str, Any]]:
        """Find resumes that performed well for similar job descriptions."""
        try:
            # Check if embedding model is available
            if self.embedding_disabled or self.embedding_model is None:
                logger.warning("Embedding model not available, returning empty results")
                return []
            
            # Embed the job description
            job_embedding = self.embedding_model.encode(job_description)
            job_vec = job_embedding.reshape(1, -1).astype('float32')
            
            # Search the index
            distances, indices = self.index.search(job_vec, k)
            
            # Get results
            results = []
            for i, idx in enumerate(indices[0]):
                if idx < 0:  # Invalid index
                    continue
                    
                # Find metadata for this index
                for item in self.meta:
                    if item["id"] == int(idx):
                        results.append({
                            "resume_path": item["resume_path"],
                            "score": item["score"],
                            "similarity": 1.0 - float(distances[0][i]),  # Convert distance to similarity
                            "timestamp": item.get("timestamp", "")
                        })
                        break
            
            return results
            
        except Exception as e:
            logger.error(f"Error finding similar resumes: {e}")
            return []


class ResumeOptimizer:
    """
    Class for optimizing resumes based on job descriptions using AI.
    """
    
    def __init__(self, llama_config=None):
        """
        Initialize the ResumeOptimizer with configuration settings.
        
        Args:
            llama_config: Configuration settings for LLM integration.
                         If None, default settings will be used.
        """
        # Try to use the provided config or create a default one
        if llama_config:
            self.llama_config = llama_config
        else:
            try:
                from job_application_automation.config.llama_config import LlamaConfig
                self.llama_config = LlamaConfig()
            except ImportError:
                logger.warning("LlamaConfig not found, using basic configuration")
                # Create a minimal compatible object with required attributes
                from types import SimpleNamespace
                self.llama_config = SimpleNamespace(
                    use_api=True, 
                    api_provider="github",
                    github_token=os.getenv("GITHUB_TOKEN", ""),
                    api_model="meta/Llama-4-Maverick-17B-128E-Instruct-FP8"
                )
                
        # Initialize components
        self.scorer = ATSScorer()
        self.llm_client = self._setup_llm_client()
        
    def _setup_llm_client(self):
        """
        Set up the LLM client based on configuration.
        
        Returns:
            LLM client instance or None if not available
        """
        # Check if we should use API-based LLM
        use_api = getattr(self.llama_config, 'use_api', True)  # Default to True if not present
        
        if not use_api:
            # Local model logic (not implemented)
            logger.warning("Local model usage not fully implemented - falling back to templates")
            return None
            
        try:
            # Determine which API provider to use
            api_provider = getattr(self.llama_config, 'api_provider', 
                                  getattr(self.llama_config, 'provider', 'openai'))
                                  
            if api_provider == "github":
                # Use GitHub token-based authentication with Azure AI SDK
                token = getattr(self.llama_config, 'github_token', os.environ.get("GITHUB_TOKEN"))
                if not token:
                    logger.warning("No GitHub token provided - LLM generation not available")
                    return None
                    
                return ChatCompletionsClient(
                    endpoint="https://models.github.ai/inference",
                    credential=AzureKeyCredential(token)
                )
                
            elif api_provider in ("groq", "openrouter", "openai"):
                # For other providers, return None for now
                # Future: implement other API clients
                logger.warning(f"{api_provider} integration not fully implemented - falling back to templates")
                return None
                
            else:
                logger.warning(f"Unknown API provider: {api_provider}")
                return None
                
        except Exception as e:
            logger.error(f"Error setting up LLM client: {e}")
            return None
    
    def get_llm_response(self, system_prompt: str, user_prompt: str) -> str:
        """
        Get a response from the LLM.
        
        Args:
            system_prompt: System message for the LLM
            user_prompt: User message for the LLM
            
        Returns:
            Generated text response
        """
        if not self.llm_client:
            logger.warning("LLM client not properly set up - cannot generate response")
            return ""
            
        try:
            # Check the structure of llm_client to handle different formats
            if isinstance(self.llm_client, dict) and "client" in self.llm_client and "model" in self.llm_client:
                # This is the format used with GitHub token
                client = self.llm_client["client"]
                model = self.llm_client["model"]
                
                from azure.ai.inference.models import SystemMessage, UserMessage
                
                logger.info(f"Sending request to LLM (model: {model})")
                response = client.complete(
                    messages=[
                        SystemMessage(system_prompt),
                        UserMessage(user_prompt),
                    ],
                    temperature=getattr(self.llama_config, 'temperature', 0.7),
                    top_p=getattr(self.llama_config, 'top_p', 0.9),
                    max_tokens=getattr(self.llama_config, 'max_tokens', 1000),
                    model=model
                )
                
                result = response.choices[0].message.content
                
            elif hasattr(self.llm_client, 'complete'):
                # Direct client with complete method
                logger.info("Sending request to direct LLM client")
                response = self.llm_client.complete(
                    system_prompt=system_prompt,
                    prompt=user_prompt,
                    temperature=getattr(self.llama_config, 'temperature', 0.7),
                    max_tokens=getattr(self.llama_config, 'max_tokens', 1000)
                )
                result = response.text if hasattr(response, 'text') else str(response)
            else:
                logger.error("Unknown LLM client format")
                return ""
            
            logger.info(f"Received {len(result)} characters from LLM")
            return result
        except Exception as e:
            logger.error(f"Error getting LLM response: {e}")
            return ""
            
    def optimize_resume(self, 
                     resume_path: str, 
                     job_description: str,
                     target_score: float = 0.8,
                     output_format: str = 'docx') -> Dict[str, Any]:
        """
        Optimize a resume for a specific job description.
        
        Args:
            resume_path: Path to resume file
            job_description: Target job description
            target_score: Target ATS score to achieve (0.0-1.0)
            output_format: Output format ('docx', 'pdf', 'txt')
            
        Returns:
            Dictionary with optimization results including path to optimized resume
        """
        logger.info(f"Starting resume optimization for {resume_path}")
        
        # Parse the resume
        resume_data = self.scorer.parse_resume(resume_path)
        if not resume_data:
            return {
                "error": "Failed to parse resume",
                "original_path": resume_path,
                "optimized_path": None,
                "original_score": 0,
                "optimized_score": 0
            }
        
        # Score original resume
        original_score = self.scorer.score_resume(resume_data, job_description)
        
        logger.info(f"Original resume score: {original_score['overall_score']}%")
        
        # Check if optimization is needed
        if original_score["overall_score"] / 100.0 >= target_score:
            logger.info(f"Resume already meets target score: {original_score['overall_score']}%")
            return {
                "original_path": resume_path,
                "optimized_path": resume_path,  # Return same path if no optimization needed
                "original_score": original_score,
                "optimized_score": original_score,
                "message": "Resume already meets target score"
            }
        
        # Extract missing keywords and suggestions
        missing_keywords = original_score["missing_important_keywords"]
        suggestions = original_score["improvement_suggestions"]
        
        # Generate optimized resume
        optimized_data = self._generate_optimized_resume(
            resume_data, 
            job_description,
            missing_keywords,
            suggestions
        )
        
        # Save optimized resume
        file_name = os.path.basename(resume_path)
        base_name, _ = os.path.splitext(file_name)
        optimized_path = str(DATA_DIR / f"{base_name}_optimized.{output_format}")
        
        # Generate the document
        self._save_optimized_resume(optimized_data, optimized_path, output_format)
        
        # Score optimized resume
        optimized_resume_data = self.scorer.parse_resume(optimized_path)
        optimized_score = self.scorer.score_resume(optimized_resume_data, job_description)
        
        logger.info(f"Optimized resume score: {optimized_score['overall_score']}%")
        
        # Return results
        return {
            "original_path": resume_path,
            "optimized_path": optimized_path,
            "original_score": original_score,
            "optimized_score": optimized_score,
            "keywords_added": [kw["keyword"] for kw in missing_keywords],
            "suggestions_implemented": suggestions
        }
    
    def _generate_optimized_resume(self, 
                               resume_data: Dict[str, Any],
                               job_description: str,
                               missing_keywords: List[Dict[str, Any]],
                               suggestions: List[str]) -> Dict[str, str]:
        """Generate optimized resume content."""
        # Convert resume data to candidate profile format
        candidate_profile = {
            "name": resume_data.get("name", ""),
            "email": resume_data.get("email", ""),
            "phone": resume_data.get("phone", ""),
            "summary": resume_data.get("summary", ""),
            "experience": resume_data.get("experience", ""),
            "education": resume_data.get("education", ""),
            "skills": resume_data.get("skills", "")
        }
        
        # Extract resume text
        resume_text = self.scorer._get_full_resume_text(resume_data)
        
        # Add missing keywords prompt enhancement
        enhancement_prompt = "Please optimize this resume with these specific improvements:\n"
        
        # Add missing keywords
        if missing_keywords:
            enhancement_prompt += "1. Include these missing keywords (naturally integrate them): "
            enhancement_prompt += ", ".join([kw["keyword"] for kw in missing_keywords])
            enhancement_prompt += "\n"
            
        # Add other suggestions
        for i, suggestion in enumerate(suggestions, start=2):
            enhancement_prompt += f"{i}. {suggestion}\n"
            
        # Generate optimized resume using the generator
        # We create a custom prompt that includes our enhancements
        prompt = f"""
        Job Description:
        {job_description}
        
        Current Resume:
        {resume_text}
        
        {enhancement_prompt}
        
        Create an optimized resume that will score well with ATS systems.
        Ensure these improvements are made while maintaining truthfulness.
        Use the same format and structure as the current resume.
        """
        
        system_prompt = "You are an expert resume optimizer that makes resumes score higher on ATS systems."
        
        # Use unified LLM client for provider agility; fallback to existing paths
        try:
            from job_application_automation.src.services.llm_client import LLMClient
            client = LLMClient(getattr(self, 'llama_config', None))
            result = client.generate(system_prompt=system_prompt, user_prompt=prompt, max_tokens=1200)
            if not result:
                raise RuntimeError("Empty LLM result")
        except Exception:
            if self.llm_client:
                result = self.get_llm_response(system_prompt, prompt)
            else:
                result = self.resume_generator._generate_text(
                    prompt=prompt,
                    system_prompt=system_prompt
                )
        
        # Parse the generated content
        optimized_data = self._parse_generated_resume(result, resume_data)
        
        return optimized_data
    
    def _parse_generated_resume(self, generated_text: str, original_data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse generated resume content into sections."""
        # Start with original data as base
        optimized_data = original_data.copy()
        
        # Try to identify sections in the generated text
        section_patterns = {
            "header": r"^(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "contact": r"(?:CONTACT|Contact)(?:\s+INFORMATION|Information)?:?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "summary": r"(?:SUMMARY|Summary|PROFESSIONAL\s+SUMMARY|Professional\s+Summary|PROFILE|Profile):?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "experience": r"(?:EXPERIENCE|Experience|WORK\s+EXPERIENCE|Work\s+Experience|EMPLOYMENT|Employment):?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "education": r"(?:EDUCATION|Education):?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "skills": r"(?:SKILLS|Skills|TECHNICAL\s+SKILLS|Technical\s+Skills):?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "projects": r"(?:PROJECTS|Projects|PROJECT\s+EXPERIENCE|Project\s+Experience):?(.*?)(?=\n\n|\n#|\n[A-Z]+:)",
            "certifications": r"(?:CERTIFICATIONS|Certifications|CERTIFICATES|Certificates):?(.*?)(?=\n\n|\n#|\n[A-Z]+:|$)"
        }
        
        # Extract sections
        for section, pattern in section_patterns.items():
            match = re.search(pattern, generated_text, re.DOTALL | re.IGNORECASE)
            if match:
                section_content = match.group(1).strip()
                if section_content:
                    optimized_data[section] = section_content
        
        return optimized_data
    
    def _save_optimized_resume(self, resume_data: Dict[str, Any], output_path: str, output_format: str) -> None:
        """Save optimized resume to file."""
        try:
            if output_format == 'docx':
                # Create Word document
                doc = Document()
                
                # Add header with name
                if "name" in resume_data:
                    doc.add_heading(resume_data["name"], level=0)
                
                # Add contact information
                if "contact" in resume_data or "email" in resume_data or "phone" in resume_data:
                    doc.add_paragraph("CONTACT INFORMATION")
                    contact_para = doc.add_paragraph()
                    if "email" in resume_data and resume_data["email"]:
                        contact_para.add_run(f"Email: {resume_data['email']} | ")
                    if "phone" in resume_data and resume_data["phone"]:
                        contact_para.add_run(f"Phone: {resume_data['phone']}")
                    if "contact" in resume_data and resume_data["contact"]:
                        contact_para.add_run(resume_data["contact"])
                
                # Add other sections
                for section in ["summary", "experience", "education", "skills", "projects", "certifications"]:
                    if section in resume_data and resume_data[section]:
                        # Add section heading
                        doc.add_heading(section.upper(), level=1)
                        # Add section content
                        lines = resume_data[section].split("\n")
                        current_para = None
                        for line in lines:
                            # Handle bullet points
                            if line.strip().startswith("-") or line.strip().startswith("•"):
                                current_para = doc.add_paragraph(style='ListBullet')
                                current_para.add_run(line.strip()[1:].strip())
                            # Handle empty lines as paragraph breaks
                            elif not line.strip() and current_para:
                                current_para = None
                            # Add to current paragraph or create new one
                            else:
                                if not current_para:
                                    current_para = doc.add_paragraph()
                                current_para.add_run(line)
                
                # Save document
                doc.save(output_path)
                
            elif output_format == 'txt':
                # Simple text format
                with open(output_path, 'w', encoding='utf-8') as f:
                    if "name" in resume_data:
                        f.write(f"{resume_data['name']}\n\n")
                    
                    # Contact information
                    f.write("CONTACT INFORMATION\n")
                    if "email" in resume_data and resume_data["email"]:
                        f.write(f"Email: {resume_data['email']}\n")
                    if "phone" in resume_data and resume_data["phone"]:
                        f.write(f"Phone: {resume_data['phone']}\n")
                    if "contact" in resume_data and resume_data["contact"]:
                        f.write(f"{resume_data['contact']}\n")
                    f.write("\n")
                    
                    # Other sections
                    for section in ["summary", "experience", "education", "skills", "projects", "certifications"]:
                        if section in resume_data and resume_data[section]:
                            f.write(f"{section.upper()}\n")
                            f.write(f"{resume_data[section]}\n\n")
            
            else:  # pdf or other formats not directly supported
                # Fall back to txt and note that conversion is needed
                self._save_optimized_resume(resume_data, output_path.replace(output_format, 'txt'), 'txt')
                logger.warning(f"Output format {output_format} not directly supported. Saved as .txt instead.")
                
            logger.info(f"Saved optimized resume to {output_path}")
                
        except Exception as e:
            logger.error(f"Error saving optimized resume: {e}")


# Example usage
if __name__ == "__main__":
    # Parse command line arguments
    import argparse
    parser = argparse.ArgumentParser(description="ATS Resume Scorer and Optimizer")
    parser.add_argument("--resume", type=str, help="Path to resume file")
    parser.add_argument("--job", type=str, help="Path to job description file")
    parser.add_argument("--score-only", action="store_true", help="Only score, don't optimize")
    parser.add_argument("--output-format", type=str, default="docx", choices=["docx", "txt", "pdf"], 
                      help="Output format for optimized resume")
    
    args = parser.parse_args()
    
    if args.resume and args.job:
        # Load job description
        with open(args.job, 'r', encoding='utf-8') as f:
            job_description = f.read()
            
        # Score and optimize resume
        scorer = ATSScorer()
        optimizer = ResumeOptimizer()
        
        if args.score_only:
            # Only score the resume
            resume_data = scorer.parse_resume(args.resume)
            score = scorer.score_resume(resume_data, job_description)
            print(f"Resume Score: {score['overall_score']}%")
            print(f"Keyword Score: {score['keyword_score']}%")
            print(f"Format Score: {score['format_score']}%")
            print(f"Semantic Score: {score['semantic_score']}%")
            print("\nMissing Keywords:")
            for kw in score['missing_important_keywords']:
                print(f"- {kw['keyword']}")
            print("\nImprovement Suggestions:")
            for suggestion in score['improvement_suggestions']:
                print(f"- {suggestion}")
        else:
            # Score and optimize
            result = optimizer.optimize_resume(
                args.resume, 
                job_description,
                output_format=args.output_format
            )
            
            print(f"Original Score: {result['original_score']['overall_score']}%")
            print(f"Optimized Score: {result['optimized_score']['overall_score']}%")
            print(f"Optimized Resume: {result['optimized_path']}")
    else:
        print("Please provide both resume and job description paths.")
        parser.print_help()