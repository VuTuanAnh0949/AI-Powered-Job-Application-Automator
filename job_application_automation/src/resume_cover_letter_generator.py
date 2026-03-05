"""
Resume and cover letter generator module using Llama 4 Mevrick.
This module provides functionality to generate personalized resumes and cover letters
based on job descriptions and user profiles.
"""

import os
import json
import logging
import asyncio
import requests
import time
from typing import Dict, Any, Optional, Union, List, Tuple
from pathlib import Path
from datetime import datetime
from enum import Enum

# Document generation libraries
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches

# Azure AI integration for GitHub token-based Llama 4 access
from azure.ai.inference import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential

# Import configuration
from job_application_automation.config.llama_config import LlamaConfig
from job_application_automation.src.utils.path_utils import get_project_root, get_data_path, get_templates_dir, get_models_dir

# Resolve project root for path references
_project_root = str(get_project_root())

# Set up logging with absolute path for the log file
log_file_path = str(get_data_path() / "resume_cover_letter_generator.log")

# Ensure log directory exists
os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Enum for cover letter templates
class CoverLetterTemplate(str, Enum):
    STANDARD = "standard"
    TECHNICAL = "technical"
    CREATIVE = "creative"
    EXECUTIVE = "executive"
    CAREER_CHANGE = "career_change"
    REFERRAL = "referral"

# Template manager for cover letters
class CoverLetterTemplateManager:
    """
    Manages different cover letter templates for various scenarios.
    """
    
    def __init__(self):
        """Initialize the template manager"""
        self.templates_dir = get_templates_dir()
        self.template_paths = {
            CoverLetterTemplate.STANDARD: self.templates_dir / "cover_letter_template.docx",
            CoverLetterTemplate.CREATIVE: self.templates_dir / "cover_letter_creative_template.docx",
            CoverLetterTemplate.TECHNICAL: self.templates_dir / "cover_letter_technical_template.docx",
            CoverLetterTemplate.EXECUTIVE: self.templates_dir / "cover_letter_executive_template.docx",
            CoverLetterTemplate.CAREER_CHANGE: self.templates_dir / "cover_letter_career_change_template.docx",
            CoverLetterTemplate.REFERRAL: self.templates_dir / "cover_letter_referral_template.docx"
        }
        
        # Template prompts for different cover letter styles
        self.template_prompts = {
            CoverLetterTemplate.STANDARD: """
            Based on the following job description and candidate resume, create a personalized cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a compelling cover letter that:
            1. Addresses the hiring manager professionally (by name if available, otherwise "Hiring Manager")
            2. Expresses enthusiasm for the position and company
            3. Highlights 2-3 most relevant experiences/skills from the resume that match the job requirements
            4. Explains why the candidate is a good fit for the company culture
            5. Includes a strong closing paragraph with a call to action
            
            The tone should be professional yet personable, confident but not arrogant.
            Format the letter properly with current date, recipient's information, and proper salutation and closing.
            """,
            
            CoverLetterTemplate.CREATIVE: """
            Based on the following job description and candidate resume, create a creative and engaging cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create an attention-grabbing cover letter that:
            1. Opens with a unique hook or story that relates to the position
            2. Conveys passion for the company's mission or industry
            3. Demonstrates creativity and innovation through specific examples
            4. Shows personality while maintaining professionalism
            5. Closes with enthusiasm and a clear call to action
            
            The tone should be engaging, conversational, and show genuine excitement.
            This template works well for creative industries (design, marketing, content creation).
            """,
            
            CoverLetterTemplate.TECHNICAL: """
            Based on the following job description and candidate resume, create a technically-focused cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a precise, technically-oriented cover letter that:
            1. Addresses the technical requirements of the role directly
            2. Highlights specific technical skills, certifications, and achievements
            3. Demonstrates problem-solving abilities through concrete examples
            4. Shows understanding of the company's technical challenges
            5. Maintains a professional, detail-oriented tone throughout
            
            Focus on quantifiable achievements and specific technical contributions.
            Include relevant technical keywords from the job description.
            This template works well for engineering, development, and IT positions.
            """,
            
            CoverLetterTemplate.EXECUTIVE: """
            Based on the following job description and candidate resume, create an executive-level cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a powerful executive cover letter that:
            1. Opens with a strong statement of value proposition
            2. Highlights strategic leadership and vision
            3. Demonstrates business impact through measurable achievements
            4. Shows industry expertise and key leadership qualities
            5. Conveys confidence and executive presence
            
            The tone should be authoritative, strategic, and refined.
            Focus on high-level achievements and leadership capabilities.
            This template works well for C-level, director, and senior management positions.
            """,
            
            CoverLetterTemplate.CAREER_CHANGE: """
            Based on the following job description and candidate resume, create a cover letter for someone changing careers.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a persuasive cover letter that:
            1. Acknowledges the career transition directly but confidently
            2. Emphasizes transferable skills that apply to the new role
            3. Explains the motivation for the career change
            4. Connects past achievements to potential future contributions
            5. Shows enthusiasm and commitment to the new field
            
            The tone should be confident and forward-looking.
            Focus on how the candidate's diverse experience offers unique value.
            This template works well for career-changers and those with non-traditional backgrounds.
            """,
            
            CoverLetterTemplate.REFERRAL: """
            Based on the following job description and candidate resume, create a cover letter that mentions a referral.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Referral Information:
            {referral_info}
            
            Create an effective referral-based cover letter that:
            1. Mentions the referrer in the opening paragraph
            2. Explains the connection to the referrer and why they recommended this role
            3. Builds on the credibility established by the referral
            4. Highlights relevant skills and experiences that match the job requirements
            5. Ends with appreciation and a clear call to action
            
            The tone should be warm yet professional.
            Leverage the referral appropriately without over-relying on it.
            This template works well when someone inside the company has referred the candidate.
            """
        }
        
    def get_template_path(self, template_type: CoverLetterTemplate) -> Path:
        """
        Get the path to a specific template file.
        
        Args:
            template_type: The type of cover letter template to use.
            
        Returns:
            Path object to the template file.
        """
        # Use the standard template if the requested one doesn't exist
        if not self.template_paths[template_type].exists():
            logger.warning(f"Template {template_type.value} not found. Using standard template.")
            return self.template_paths[CoverLetterTemplate.STANDARD]
        
        return self.template_paths[template_type]
    
    def get_template_prompt(self, template_type: CoverLetterTemplate) -> str:
        """
        Get the prompt for a specific template type.
        
        Args:
            template_type: The type of cover letter template to use.
            
        Returns:
            String prompt for the selected template.
        """
        return self.template_prompts.get(template_type, self.template_prompts[CoverLetterTemplate.STANDARD])
    
    def select_best_template(self, job_description: str, candidate_profile: Dict[str, Any]) -> CoverLetterTemplate:
        """
        Automatically select the most appropriate template based on the job and candidate.
        
        Args:
            job_description: The job description text.
            candidate_profile: The candidate's profile information.
            
        Returns:
            The most appropriate CoverLetterTemplate type.
        """
        job_desc_lower = job_description.lower()
        
        # Check for executive positions
        if any(term in job_desc_lower for term in ['ceo', 'cto', 'cfo', 'chief', 'director', 'vp', 'vice president', 'head of']):
            return CoverLetterTemplate.EXECUTIVE
            
        # Check for technical positions
        if any(term in job_desc_lower for term in ['engineer', 'developer', 'programmer', 'architect', 'data scientist']):
            return CoverLetterTemplate.TECHNICAL
            
        # Check for creative positions
        if any(term in job_desc_lower for term in ['design', 'creative', 'writer', 'artist', 'content', 'marketing']):
            return CoverLetterTemplate.CREATIVE
            
        # Default to standard
        return CoverLetterTemplate.STANDARD


class ResumeGenerator:
    """
    Class for generating personalized resumes and cover letters using LLM.
    """

    def __init__(self, config: Optional[LlamaConfig] = None):
        """
        Initialize the ResumeGenerator with configuration settings.
        
        Args:
            config: Configuration settings for LLM integration.
                   If None, default settings will be used.
        """
        self.config = config or LlamaConfig()
        self._setup_llm()
        self.template_manager = CoverLetterTemplateManager()
        self.llm_available = False
        self.api_available = False
        self.github_client = None
        
    def _setup_llm(self) -> None:
        """Set up the LLM for resume and cover letter generation."""
        try:
            # Create models directory if it doesn't exist 
            models_dir = os.path.join(_project_root, "models")
            os.makedirs(models_dir, exist_ok=True)

            # Create data directories if they don't exist
            output_dir = os.path.join(_project_root, "data", "generated_cover_letters")
            os.makedirs(output_dir, exist_ok=True)
            
            # First check if API is available
            if hasattr(self.config, 'use_api') and self.config.use_api:
                api_config = self.config.get_api_config()
                if api_config:
                    if self.config.api_provider == "github":
                        # Set up GitHub token-based Llama 4 client
                        self._setup_github_llama4_client(api_config)
                    else:
                        self.api_available = self._test_api_connection(api_config)
                        
                    if self.api_available or self.github_client:
                        logger.info(f"Successfully connected to {self.config.api_provider} API")
                        self.llm_available = False  # Use API instead of local model
                        return
                    else:
                        logger.warning(f"Failed to connect to {self.config.api_provider} API, falling back to local model")
                
            # Check if model file exists
            model_paths = [
                os.path.join(_project_root, "models", "llama-4-mevrick"),
                os.path.join(_project_root, "models", "llama-4-mevrick.gguf"),
                os.path.join(_project_root, "models", "llama-3-8b.gguf"),
                os.path.join(_project_root, "models", "llama-2-7b.gguf"),
                os.path.join(os.path.expanduser("~"), ".cache", "models", "llama-4-mevrick.gguf"),
                os.path.join(os.path.expanduser("~"), ".cache", "models", "llama-3-8b.gguf")
            ]
            
            model_path = None
            for path in model_paths:
                if os.path.exists(path):
                    model_path = path
                    logger.info(f"Found model file at {model_path}")
                    break
                    
            if not model_path:
                logger.warning("Model file not found - will use template-based generation instead")
                self.llm_available = False
                return
                
            try:
                # Try to import the llama_cpp module
                try:
                    import llama_cpp
                    logger.info("Successfully imported llama_cpp module")
                except ImportError as imp_err:
                    logger.warning(f"Failed to import llama_cpp: {imp_err} - using template-based generation")
                    self.llm_available = False
                    return
                    
                # Try to load the model
                logger.info(f"Loading model from {model_path} with {self.config.n_threads} threads and {self.config.n_gpu_layers} GPU layers")
                self.llm = llama_cpp.Llama(
                    model_path=model_path,
                    n_ctx=self.config.context_length,
                    n_threads=self.config.n_threads,
                    n_gpu_layers=self.config.n_gpu_layers
                )
                self.llm_available = True
                logger.info("LLM initialized successfully")
                
            except ImportError as e:
                logger.warning(f"llama_cpp package not installed: {e} - falling back to template-based generation")
                self.llm_available = False
                
            except Exception as e:
                logger.warning(f"Error loading model: {e} - falling back to template-based generation")
                self.llm_available = False
                
        except Exception as e:
            logger.warning(f"Error setting up LLM: {e} - will use template-based generation instead")
            self.llm_available = False
    
    def _setup_github_llama4_client(self, api_config: Dict[str, Any]) -> None:
        """Set up GitHub token-based Llama 4 client."""
        try:
            endpoint = api_config.get("endpoint", "https://models.github.ai/inference")
            token = api_config.get("token")
            model = api_config.get("model", "meta/Llama-4-Maverick-17B-128E-Instruct-FP8")
            
            if not token:
                logger.error("GitHub token not found. Unable to set up GitHub-based Llama 4 client.")
                return
                
            self.github_client = ChatCompletionsClient(
                endpoint=endpoint,
                credential=AzureKeyCredential(token),
            )
            
            logger.info(f"GitHub-based Llama 4 client initialized successfully with model {model}")
            self.api_available = True
        except Exception as e:
            logger.error(f"Error setting up GitHub Llama 4 client: {e}")
            self.github_client = None
    
    def _test_api_connection(self, api_config: Dict[str, Any]) -> bool:
        """Test connection to the API provider."""
        try:
            api_base = api_config.get("api_base", "")
            api_key = api_config.get("api_key", "")
            model = api_config.get("model", "")
            timeout = api_config.get("timeout", 10)
            
            if not api_base or not api_key or not model:
                logger.error("Missing API configuration parameters")
                return False
            
            # Prepare a minimal test request
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # Different request format based on provider
            if "openrouter" in api_base:
                # OpenRouter endpoint test
                endpoint = f"{api_base.rstrip('/')}/models"
                response = requests.get(
                    endpoint,
                    headers=headers,
                    timeout=timeout
                )
            else:
                # Groq or OpenAI-compatible endpoint test
                endpoint = f"{api_base.rstrip('/')}/models"
                response = requests.get(
                    endpoint,
                    headers=headers,
                    timeout=timeout
                )
            
            if response.status_code == 200:
                return True
            else:
                logger.error(f"API test failed with status code: {response.status_code}, message: {response.text}")
                return False
                
        except Exception as e:
            logger.error(f"API connection test failed: {e}")
            return False
    
    def _generate_text_with_api(self, prompt: str) -> str:
        """Generate text using the API."""
        api_config = self.config.get_api_config()
        
        if not api_config:
            logger.error("API config not available for text generation")
            return "API error: Could not generate text."
            
        # If using GitHub-based Llama 4
        if self.github_client and self.config.api_provider == "github":
            try:
                model = api_config.get("model", "meta/Llama-4-Maverick-17B-128E-Instruct-FP8")
                
                # Prepare the messages for chat completion
                messages = [
                    SystemMessage("You are a helpful assistant specialized in writing professional resumes and cover letters."),
                    UserMessage(prompt),
                ]
                
                # Call the GitHub-hosted Llama 4 model
                response = self.github_client.complete(
                    messages=messages,
                    temperature=self.config.temperature,
                    top_p=self.config.top_p,
                    max_tokens=min(4000, self.config.context_length),
                    model=model
                )
                
                generated_text = response.choices[0].message.content
                logger.info("Successfully generated text using GitHub-based Llama 4")
                return generated_text
                
            except Exception as e:
                logger.error(f"Error generating text with GitHub-based Llama 4: {e}")
                return f"GitHub API error: {str(e)}"
        
        # For other API providers (existing code)
        if not self.api_available:
            logger.error("API not available for text generation")
            return "API error: Could not generate text."
            
        try:
            api_base = api_config.get("api_base", "")
            api_key = api_config.get("api_key", "")
            model = api_config.get("model", "")
            timeout = api_config.get("timeout", 60)
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # Add http-referer header for OpenRouter
            if "openrouter" in api_base:
                headers["HTTP-Referer"] = "https://github.com/job-application-automation"
                headers["X-Title"] = "Job Application Automation Tool"
            
            # Prepare the request payload
            payload = {
                "model": model,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": self.config.temperature,
                "top_p": self.config.top_p,
                "max_tokens": min(4000, self.config.context_length)
            }
            
            # Send the request
            endpoint = f"{api_base.rstrip('/')}/chat/completions"
            logger.info(f"Sending request to {endpoint} with model {model}")
            
            response = requests.post(
                endpoint,
                headers=headers,
                json=payload,
                timeout=timeout
            )
            
            # Process the response
            if response.status_code == 200:
                response_json = response.json()
                if self.config.api_provider == "openrouter":
                    generated_text = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
                    # Log usage information if available
                    if "usage" in response_json:
                        usage = response_json["usage"]
                        logger.info(f"API usage: prompt_tokens={usage.get('prompt_tokens')}, completion_tokens={usage.get('completion_tokens')}")
                else:
                    # Groq or other OpenAI-compatible APIs
                    generated_text = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
                    # Log usage information if available
                    if "usage" in response_json:
                        usage = response_json["usage"]
                        logger.info(f"API usage: prompt_tokens={usage.get('prompt_tokens')}, completion_tokens={usage.get('completion_tokens')}")
                
                return generated_text
            else:
                logger.error(f"API request failed with status code: {response.status_code}, message: {response.text}")
                return f"API error: {response.status_code} - Could not generate text."
                
        except Exception as e:
            logger.error(f"Error generating text with API: {e}")
            return f"API error: {str(e)}"
            
    def _generate_text_with_llm(self, prompt: str) -> str:
        """Generate text using the local LLM."""
        if not self.llm_available:
            logger.warning("LLM not available for text generation")
            # Use template-based fallback
            return self._template_based_fallback(prompt)
            
        try:
            # Process with local Llama model
            if self.llm:
                output = self.llm(
                    prompt,
                    max_tokens=self.config.context_length,
                    temperature=self.config.temperature,
                    top_p=self.config.top_p,
                    echo=False
                )
                generated_text = output.get("choices", [{}])[0].get("text", "").strip()
                return generated_text
            else:
                return self._template_based_fallback(prompt)
                
        except Exception as e:
            logger.error(f"Error generating with local LLM: {e}")
            return self._template_based_fallback(prompt)
            
    def _template_based_fallback(self, prompt: str) -> str:
        """Generate text based on templates when LLM is not available."""
        # Simple template-based generation logic here
        logger.warning("Using template-based fallback for text generation")
        
        # Identify the type of content needed based on prompt keywords
        if "resume" in prompt.lower() or "skills section" in prompt.lower():
            return "Skills:\n- Technical skills relevant to the position\n- Soft skills like communication and teamwork\n- Industry-specific knowledge\n- Tools and methodologies"
        elif "cover letter" in prompt.lower():
            return "Dear Hiring Manager,\n\nI am writing to express my interest in the position. With my background and experience, I believe I would be a valuable addition to your team.\n\nThank you for your consideration.\n\nSincerely,\n[Your Name]"
        else:
            return "Generated content based on the provided information."
                
    def generate_text(self, prompt: str) -> str:
        """Generate text using available methods (API or local LLM)."""
        # Choose the appropriate generation method
        if self.config.use_api and self.api_available:
            logger.info(f"Generating text with {self.config.api_provider} API")
            return self._generate_text_with_api(prompt)
        elif self.llm_available:
            logger.info("Generating text with local LLM")
            return self._generate_text_with_llm(prompt)
        else:
            logger.warning("No text generation method available, using fallback")
            return self._template_based_fallback(prompt)

    # The remaining methods can use self.generate_text instead of directly using self.llm
    
    def generate_resume(self, 
                  job_description: str, 
                  candidate_profile: Dict[str, Any],
                  output_path: Optional[str] = None) -> Tuple[str, str]:
        """
        Generate a personalized resume based on a job description and candidate profile.
        
        Args:
            job_description: The job description to tailor the resume to.
            candidate_profile: The candidate's profile data.
            output_path: Optional path to save the generated resume.
                         If None, a default path will be used.
                         
        Returns:
            A tuple containing (file_path, resume_content).
        """
        try:
            prompt = self._prepare_resume_prompt(job_description, candidate_profile)
            
            # Generate resume content using unified client for provider agility
            try:
                from job_application_automation.src.services.llm_client import LLMClient
                client = LLMClient(self.config)
                resume_content = client.generate(
                    system_prompt="You are a helpful assistant specialized in writing professional resumes.",
                    user_prompt=prompt,
                    max_tokens=min(4000, self.config.context_length),
                ) or self._template_based_fallback(prompt)
            except Exception:
                resume_content = self.generate_text(prompt)
            
            # Compute default output path
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                job_title = candidate_profile.get("target_job_title", "resume")
                output_dir = os.path.join(_project_root, "data", "generated_resumes")
                os.makedirs(output_dir, exist_ok=True)
                output_path = os.path.join(output_dir, f"{timestamp}_{job_title.replace(' ', '_')}.docx")

            # Create a proper document (template if available)
            final_path = self._create_resume_document(resume_content, output_path)
            logger.info(f"Generated resume saved to {final_path}")
            return final_path, resume_content
            
        except Exception as e:
            logger.error(f"Error generating resume: {e}")
            # Return a minimal result in case of failure
            return "", "Error generating resume."
            
    def generate_cover_letter(self, 
                       job_description: str, 
                       candidate_resume: str,
                       company_info: str,
                       output_path: Optional[str] = None,
                       template_type: Optional[CoverLetterTemplate] = None,
                       referral_info: Optional[str] = None) -> Tuple[str, str]:
        """
        Generate a personalized cover letter based on job description and resume.
        
        Args:
            job_description: The job description to tailor the cover letter to.
            candidate_resume: The candidate's resume content.
            company_info: Information about the company.
            output_path: Optional path to save the generated cover letter.
                        If None, a default path will be used.
            template_type: Optional template type for the cover letter.
                          If None, a standard template will be used.
            referral_info: Optional referral information to include.
                          
        Returns:
            A tuple containing (file_path, cover_letter_content).
        """
        try:
            prompt = self._prepare_cover_letter_prompt(
                job_description, 
                candidate_resume, 
                company_info,
                template_type or CoverLetterTemplate.STANDARD,
                referral_info
            )
            
            # Generate cover letter content using unified client for provider agility
            try:
                from job_application_automation.src.services.llm_client import LLMClient
                client = LLMClient(self.config)
                cover_letter_content = client.generate(
                    system_prompt="You are a helpful assistant specialized in writing professional cover letters.",
                    user_prompt=prompt,
                    max_tokens=min(3000, self.config.context_length),
                ) or self._template_based_fallback(prompt)
            except Exception:
                cover_letter_content = self.generate_text(prompt)
            
            # Compute default output path
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                job_title = "cover_letter"
                output_dir = os.path.join(_project_root, "data", "generated_cover_letters")
                os.makedirs(output_dir, exist_ok=True)
                output_path = os.path.join(output_dir, f"{timestamp}_{job_title.replace(' ', '_')}.docx")

            # Create a proper document using template when available
            template_path = self.template_manager.get_template_path(template_type)
            final_path = self._create_cover_letter_document(cover_letter_content, str(template_path), output_path)
            logger.info(f"Generated cover letter saved to {final_path}")
            return final_path, cover_letter_content
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            # Return a minimal result in case of failure
            return "", "Error generating cover letter."
    
    def _prepare_resume_prompt(self, 
                        job_description: str, 
                        candidate_profile: Dict[str, Any]) -> str:
        """
        Prepare the prompt for resume generation with enhanced AI analysis.
        """
        # Extract key requirements from job description using AI
        key_requirements = self._extract_job_requirements(job_description)
        
        # Score and sort candidate skills based on job requirements
        scored_skills = self._score_skills_match(candidate_profile.get("skills", []), key_requirements)
        
        # Format candidate profile with prioritized skills
        profile_sections = []
        for section, content in candidate_profile.items():
            if section == "skills":
                # Order skills by relevance score
                skills_text = "\nSkills:\n" + "\n".join(f"- {skill}" for skill, _ in scored_skills)
                profile_sections.append(skills_text)
            else:
                profile_sections.append(f"{section}:\n{content}")
                
        profile_text = "\n\n".join(profile_sections)
        
        # Generate focused prompt
        prompt = self.config.resume_template_prompt.format(
            job_description=job_description,
            candidate_profile=profile_text,
            key_requirements="\n".join(f"- {req}" for req in key_requirements)
        )
        
        return prompt

    def _extract_job_requirements(self, job_description: str) -> List[str]:
        """Extract key requirements from job description using AI analysis."""
        try:
            # Use the LLM to analyze requirements
            analysis_prompt = """
            Analyze the following job description and extract the key requirements including:
            - Technical skills
            - Soft skills
            - Experience level
            - Education requirements
            - Industry knowledge
            
            Job Description:
            {description}
            
            Format each requirement as a concise bullet point.
            """
            
            response = self._generate_text(
                prompt=analysis_prompt.format(description=job_description),
                system_prompt="You are an expert job requirements analyzer."
            )
            
            # Parse bullet points from response
            requirements = [line.strip("- ").strip() 
                          for line in response.split("\n") 
                          if line.strip().startswith("-")]
            
            return requirements
            
        except Exception as e:
            logger.error(f"Error extracting job requirements: {e}")
            return []

    def _score_skills_match(self, 
                      candidate_skills: List[str], 
                      job_requirements: List[str]) -> List[Tuple[str, float]]:
        """Score candidate skills against job requirements using AI matching."""
        try:
            # Prepare scoring prompt
            scoring_prompt = """
            Score how well each candidate skill matches the job requirements.
            Consider direct matches, related skills, and transferable skills.
            
            Job Requirements:
            {requirements}
            
            Candidate Skills:
            {skills}
            
            For each skill, respond with: skill|score
            Score from 0.0 to 1.0 where 1.0 is a perfect match.
            """
            
            # Format the prompt
            formatted_prompt = scoring_prompt.format(
                requirements="\n".join(f"- {req}" for req in job_requirements),
                skills="\n".join(f"- {skill}" for skill in candidate_skills)
            )
            
            # If we have AI available, use it
            if self.llm_available or self.api_available:
                response = self.generate_text(formatted_prompt)
                
                # Parse response
                scored_skills = []
                for line in response.split("\n"):
                    if "|" in line:
                        parts = line.split("|")
                        if len(parts) == 2:
                            skill = parts[0].strip()
                            try:
                                score = float(parts[1].strip())
                                score = max(0.0, min(1.0, score))  # Clamp to 0.0-1.0
                                scored_skills.append((skill, score))
                            except ValueError:
                                # Skip malformed lines
                                continue
                
                # If we got results, return them
                if scored_skills:
                    return sorted(scored_skills, key=lambda x: x[1], reverse=True)
            
            # Fallback: simple text matching algorithm
            scored_skills = []
            for skill in candidate_skills:
                skill_lower = skill.lower()
                
                # Calculate best match score against requirements
                best_score = 0.0
                for req in job_requirements:
                    req_lower = req.lower()
                    
                    # Direct match
                    if skill_lower == req_lower or skill_lower in req_lower:
                        score = 1.0
                    # Partial match (skill words in requirement)
                    elif any(word in req_lower for word in skill_lower.split()):
                        score = 0.7
                    # Related field match
                    elif self._are_related_tech(skill_lower, req_lower):
                        score = 0.5
                    # No match
                    else:
                        score = 0.1
                        
                    best_score = max(best_score, score)
                
                scored_skills.append((skill, best_score))
                
            return sorted(scored_skills, key=lambda x: x[1], reverse=True)
            
        except Exception as e:
            logger.error(f"Error scoring skills match: {e}")
            # Return unsorted skills with default score
            return [(skill, 0.5) for skill in candidate_skills]
            
    def _are_related_tech(self, skill: str, requirement: str) -> bool:
        """Check if a skill and requirement are related in the tech domain."""
        # Define related technology groups
        tech_groups = [
            {"python", "django", "flask", "fastapi", "tornado", "pyramid", "web2py", "bottle", "scipy", "numpy", "pandas"},
            {"javascript", "typescript", "node", "vue", "react", "angular", "express", "next.js", "nuxt", "svelte"},
            {"java", "spring", "hibernate", "jpa", "jakarta", "j2ee", "servlet", "struts", "vaadin"},
            {"cloud", "aws", "azure", "gcp", "google cloud", "serverless", "kubernetes", "docker", "containers"},
            {"data", "sql", "mysql", "postgresql", "mongodb", "nosql", "database", "analytics", "warehouse", "big data"},
            {"devops", "ci/cd", "jenkins", "github actions", "gitlab ci", "terraform", "ansible", "puppet", "chef"},
            {"machine learning", "ml", "ai", "deep learning", "tensorflow", "pytorch", "keras", "scikit-learn"}
        ]
        
        # Check if skill and requirement are in the same group
        for group in tech_groups:
            if any(term in skill for term in group) and any(term in requirement for term in group):
                return True
                
        return False
            
    def _prepare_cover_letter_prompt(self,
                              job_description: str,
                              candidate_resume: str,
                              company_info: str,
                              template_type: CoverLetterTemplate,
                              referral_info: Optional[str] = None) -> str:
        """
        Prepare the prompt for cover letter generation with enhanced AI analysis.
        """
        # Get the template prompt for the specified type
        template_prompt = self.template_manager.get_template_prompt(template_type)
        
        # Format the template prompt with our data
        prompt_data = {
            "job_description": job_description,
            "candidate_resume": candidate_resume,
            "company_info": company_info
        }
        
        # Add referral info if provided
        if referral_info:
            prompt_data["referral_info"] = referral_info
        
        # Format the prompt
        try:
            formatted_prompt = template_prompt.format(**prompt_data)
        except KeyError as e:
            # Handle case where template needs a field we don't have
            logger.warning(f"Template missing required field: {e}")
            # Fall back to standard template
            standard_template = self.template_manager.get_template_prompt(CoverLetterTemplate.STANDARD)
            formatted_prompt = standard_template.format(**prompt_data)
            
        # Add system instruction for high quality cover letter
        system_instruction = "Create a professional, tailored cover letter that matches the candidate's qualifications to the job requirements. Use a confident tone, avoid clichés, and focus on specific achievements relevant to the role."
        
        # Combine system instruction and prompt
        final_prompt = f"{system_instruction}\n\n{formatted_prompt}"
        
        return final_prompt
            
    def _generate_text(self, 
                prompt: str, 
                system_prompt: Optional[str] = None, 
                **kwargs) -> str:
        """
        Generate text using available generation method (API or local LLM).
        
        Args:
            prompt: The prompt to generate text from.
            system_prompt: Optional system prompt for guidance.
            **kwargs: Additional parameters for the text generation.
            
        Returns:
            Generated text as a string.
        """
        try:
            # Combine prompts if system prompt is provided
            combined_prompt = prompt
            if system_prompt:
                combined_prompt = f"{system_prompt}\n\n{prompt}"
                
            # Choose generation method
            if self.config.use_api and self.api_available:
                return self._generate_text_with_api(combined_prompt)
            elif self.llm_available:
                return self._generate_text_with_llm(combined_prompt)
            else:
                return self._template_based_fallback(combined_prompt)
                
        except Exception as e:
            logger.error(f"Error in _generate_text: {e}")
            return self._template_based_fallback(prompt)
            
    def _create_resume_document(self, content: str, output_path: str) -> str:
        """
        Create a properly formatted resume document from generated content.
        
        Args:
            content: The resume content text.
            output_path: Path to save the document to.
            
        Returns:
            Path to the created document.
        """
        try:
            # Create a new document or use template
            template_path = os.path.join(_project_root, "templates", "resume_template.docx")
            
            # If template exists, use it
            if os.path.exists(template_path):
                doc = DocxTemplate(template_path)
                
                # Extract sections from content
                context = self._parse_resume_sections(content)
                
                # Render template with context
                doc.render(context)
                doc.save(output_path)
            else:
                # No template, create a simple document
                doc = Document()
                
                # Add title
                title = doc.add_heading("Resume", 0)
                
                # Add content by paragraphs
                for paragraph in content.split("\n\n"):
                    if paragraph.strip():
                        if paragraph.strip().endswith(":"):
                            # This is likely a section header
                            doc.add_heading(paragraph.strip(), level=1)
                        else:
                            p = doc.add_paragraph(paragraph.strip())
                
                # Save document
                doc.save(output_path)
                
            logger.info(f"Created resume document at {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating resume document: {e}")
            
            # Fallback - save as plain text
            text_path = output_path.replace(".docx", ".txt")
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Saved resume as plain text at {text_path}")
            return text_path

    def _create_cover_letter_document(self, content: str, template_path: str, output_path: str) -> str:
        """
        Create a properly formatted cover letter document from generated content.
        Uses a template if available, otherwise creates a simple DOCX.
        """
        try:
            # If template exists and is a docx template, try to render with a simple context
            if os.path.exists(template_path):
                try:
                    doc = DocxTemplate(template_path)
                    context = {"content": content}
                    doc.render(context)
                    doc.save(output_path)
                    logger.info(f"Rendered cover letter using template at {output_path}")
                    return output_path
                except Exception as e:
                    logger.warning(f"Failed to render cover letter with template: {e}. Falling back to plain document.")

            # Fallback: write a simple DOCX with paragraphs
            doc = Document()
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.save(output_path)
            return output_path
        except Exception as e:
            logger.error(f"Error creating cover letter document: {e}")
            text_path = output_path.replace(".docx", ".txt")
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(content)
            return text_path
            
    def _parse_resume_sections(self, content: str) -> Dict[str, Any]:
        """Parse resume content into sections for template rendering."""
        context = {
            "name": "",
            "email": "",
            "phone": "",
            "summary": "",
            "skills": [],
            "experience": [],
            "education": []
        }
        
        # Simple section-based parsing
        current_section = None
        current_content = []
        
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
                
            # Check for section headers
            lower_line = line.lower()
            if lower_line.endswith(":"):
                # Save previous section
                if current_section and current_content:
                    if current_section == "skills":
                        # Parse skills as list
                        skills = []
                        for skill_line in current_content:
                            if skill_line.startswith("- "):
                                skills.append(skill_line[2:])
                            else:
                                skills.extend([s.strip() for s in skill_line.split(",")])
                        context["skills"] = skills
                    elif current_section in context:
                        context[current_section] = "\n".join(current_content)
                
                # Start new section
                section_name = lower_line[:-1]  # Remove colon
                if "name" in section_name:
                    current_section = "name"
                elif "contact" in section_name:
                    current_section = "contact"
                elif "summary" in section_name or "objective" in section_name:
                    current_section = "summary"
                elif "skill" in section_name:
                    current_section = "skills"
                elif "experience" in section_name or "work" in section_name:
                    current_section = "experience"
                elif "education" in section_name:
                    current_section = "education"
                else:
                    current_section = section_name
                
                current_content = []
            else:
                # Add line to current section
                if current_section == "contact":
                    # Try to extract contact details
                    if "@" in line:
                        context["email"] = line.strip()
                    elif any(c.isdigit() for c in line):
                        context["phone"] = line.strip()
                    else:
                        current_content.append(line)
                elif current_section == "name":
                    context["name"] = line.strip()
                else:
                    current_content.append(line)
        
        # Handle last section
        if current_section and current_content:
            if current_section == "skills":
                skills = []
                for skill_line in current_content:
                    if skill_line.startswith("- "):
                        skills.append(skill_line[2:])
                    else:
                        skills.extend([s.strip() for s in skill_line.split(",")])
                context["skills"] = skills
            elif current_section in context:
                context[current_section] = "\n".join(current_content)
                
        return context
    
    def _extract_profile_from_resume(self, resume_content: str) -> Dict[str, Any]:
        """
        Extract a profile dictionary from resume content.
        
        Args:
            resume_content: The resume text content.
            
        Returns:
            A dictionary with profile information.
        """
        profile = {}
        
        try:
            # Simple extraction of main sections
            sections = self._parse_content_to_sections(resume_content)
            
            # Extract name - usually the first heading or first line
            lines = resume_content.split('\n')
            for line in lines[:5]:
                if line.strip() and not line.startswith('#'):
                    profile['name'] = line.strip()
                    break
            
            # Extract other sections
            if 'contact_information' in sections:
                profile['contact'] = sections['contact_information']
            
            if 'skills' in sections:
                profile['skills'] = sections['skills'].split('\n')
            
            if 'experience' in sections:
                profile['experience'] = sections['experience']
            
            if 'education' in sections:
                profile['education'] = sections['education']
                
        except Exception as e:
            logger.error(f"Error extracting profile from resume: {e}")
            
        return profile
    
    def _extract_company_name(self, job_description: str) -> str:
        """
        Extract company name from job description.
        
        Args:
            job_description: Job description text.
            
        Returns:
            Extracted company name or "Unknown".
        """
        try:
            # A very simple heuristic - look for common patterns
            patterns = [
                r"(?i)at\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+is\s+looking)",
                r"(?i)with\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+as)",
                r"(?i)join\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+as)",
                r"(?i)about\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s*:)"
            ]
            
            import re
            for pattern in patterns:
                matches = re.search(pattern, job_description)
                if matches:
                    return matches.group(1).strip()
                    
            # Default if no match found
            return "Unknown"
            
        except Exception as e:
            logger.error(f"Error extracting company name: {e}")
            return "Unknown"
            
    def _extract_job_title(self, job_description: str) -> str:
        """
        Extract job title from job description.
        
        Args:
            job_description: Job description text.
            
        Returns:
            Extracted job title or "Position".
        """
        try:
            # Simple heuristic to find job title
            patterns = [
                r"(?i)(Software Engineer|Data Scientist|Machine Learning Engineer|Full Stack Developer|Frontend Developer|Backend Developer|DevOps Engineer|Cloud Engineer|AI Engineer|ML Engineer)(?=\s+at)",
                r"(?i)(Software Engineer|Data Scientist|Machine Learning Engineer|Full Stack Developer|Frontend Developer|Backend Developer|DevOps Engineer|Cloud Engineer|AI Engineer|ML Engineer)(?=\s+position)"
            ]
            
            import re
            for pattern in patterns:
                matches = re.search(pattern, job_description)
                if matches:
                    return matches.group(1).strip()
                    
            # Look at the beginning of the job description
            first_line = job_description.split('\n')[0]
            if len(first_line) < 100:  # Usually the title is short
                return first_line.strip()
                
            # Default if no match found
            return "Position"
            
        except Exception as e:
            logger.error(f"Error extracting job title: {e}")
            return "Position"


# Example usage
def main():
    # Example job description
    job_description = """
    Software Engineer at TechCorp
    
    We're looking for a skilled Software Engineer to join our team at TechCorp. 
    The ideal candidate will have experience with Python, JavaScript, and cloud technologies.
    
    Responsibilities:
    - Develop and maintain web applications
    - Write clean, efficient, and maintainable code
    - Collaborate with cross-functional teams
    
    Requirements:
    - Bachelor's degree in Computer Science or related field
    - 2+ years of experience in software development
    - Proficiency in Python and JavaScript
    - Experience with cloud platforms (AWS, Azure, or GCP)
    - Strong problem-solving skills
    """
    
    # Example candidate profile
    candidate_profile = {
        "name": "Jane Doe",
        "email": "jane.doe@email.com",
        "phone": "555-123-4567",
        "summary": "Software Engineer with 3 years of experience in Python and JavaScript development.",
        "skills": ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker"],
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Tech Solutions Inc.",
                "duration": "2021-2023",
                "description": "Developed and maintained web applications using Python and JavaScript."
            },
            {
                "title": "Junior Developer",
                "company": "StartUp Co.",
                "duration": "2020-2021",
                "description": "Assisted in building frontend components with React."
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "institution": "State University",
                "year": "2020"
            }
        ]
    }
    
    # Example company info
    company_info = """
    TechCorp is a leading technology company specializing in cloud solutions and web applications.
    Founded in 2010, the company has grown to over 500 employees and serves clients worldwide.
    The company culture emphasizes innovation, collaboration, and work-life balance.
    """
    
    # Create resume generator
    resume_generator = ResumeGenerator()
    
    # Generate resume
    resume_path, resume_content = resume_generator.generate_resume(
        job_description=job_description,
        candidate_profile=candidate_profile
    )
    
    print(f"Resume saved to: {resume_path}")
    
    # Generate standard cover letter
    cover_letter_path, cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.STANDARD
    )
    
    print(f"Cover letter saved to: {cover_letter_path}")
    
    # Generate technical cover letter
    tech_cover_letter_path, tech_cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.TECHNICAL
    )
    
    print(f"Technical cover letter saved to: {tech_cover_letter_path}")
    
    # Generate referral cover letter with referral info
    referral_info = "John Smith, Senior Engineer at TechCorp, who I worked with at my previous company."
    referral_cover_letter_path, referral_cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.REFERRAL,
        referral_info=referral_info
    )
    
    print(f"Referral cover letter saved to: {referral_cover_letter_path}")


if __name__ == "__main__":
    main()